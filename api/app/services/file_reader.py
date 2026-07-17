"""Extract plain text from common document types for journal import.

Supported: .txt, .md, .markdown, .csv, .html/.htm, .rtf, .pdf, .docx.
Binary formats are parsed with dedicated libraries; text formats are decoded
with a UTF-8 → Latin-1 fallback. Content is truncated to the entry max length.
"""

from __future__ import annotations

import csv
import io
import logging
import re
from dataclasses import dataclass
from html.parser import HTMLParser
from pathlib import PurePosixPath
from typing import FrozenSet

logger = logging.getLogger(__name__)

# Keep in sync with entry schema max content length.
MAX_EXTRACTED_CHARS = 50_000
MAX_UPLOAD_BYTES = 10 * 1024 * 1024  # 10 MB

ALLOWED_EXTENSIONS: FrozenSet[str] = frozenset(
    {
        ".txt",
        ".md",
        ".markdown",
        ".csv",
        ".html",
        ".htm",
        ".rtf",
        ".pdf",
        ".docx",
    }
)

# Extension → accepted MIME types (browser Content-Type is advisory only).
ALLOWED_MIME_TYPES: dict[str, FrozenSet[str]] = {
    ".txt": frozenset({"text/plain", "application/octet-stream"}),
    ".md": frozenset({"text/plain", "text/markdown", "application/octet-stream"}),
    ".markdown": frozenset({"text/plain", "text/markdown", "application/octet-stream"}),
    ".csv": frozenset({"text/csv", "text/plain", "application/csv", "application/octet-stream"}),
    ".html": frozenset({"text/html", "application/octet-stream"}),
    ".htm": frozenset({"text/html", "application/octet-stream"}),
    ".rtf": frozenset({"application/rtf", "text/rtf", "application/octet-stream"}),
    ".pdf": frozenset({"application/pdf", "application/octet-stream"}),
    ".docx": frozenset(
        {
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/octet-stream",
            "application/zip",
        }
    ),
}


class FileReaderError(Exception):
    """Raised when a file cannot be accepted or text cannot be extracted."""

    def __init__(self, message: str, *, status_code: int = 400):
        super().__init__(message)
        self.status_code = status_code


@dataclass(frozen=True)
class ExtractedDocument:
    text: str
    filename: str
    extension: str
    mime_type: str
    truncated: bool


def sanitize_filename(filename: str | None) -> str:
    """Return a basename-only, traversal-safe filename."""
    if not filename or not filename.strip():
        raise FileReaderError("Filename is required")

    # Strip any directory components the client may have sent.
    name = PurePosixPath(filename.replace("\\", "/")).name
    name = name.strip().lstrip(".")
    # Drop control chars and path separators that survived.
    name = re.sub(r'[\x00-\x1f<>:"/\\|?*]', "_", name)
    if not name or name in {".", ".."}:
        raise FileReaderError("Invalid filename")
    # Cap length while preserving extension when possible.
    if len(name) > 200:
        stem = PurePosixPath(name).stem[:180]
        suffix = PurePosixPath(name).suffix[:20]
        name = f"{stem}{suffix}"
    return name


def resolve_extension(filename: str) -> str:
    ext = PurePosixPath(filename).suffix.lower()
    if ext not in ALLOWED_EXTENSIONS:
        allowed = ", ".join(sorted(ALLOWED_EXTENSIONS))
        raise FileReaderError(
            f"Unsupported file type '{ext or '(none)'}'. Allowed: {allowed}"
        )
    return ext


def validate_mime(extension: str, mime_type: str | None) -> str:
    """Accept a MIME type if it matches the extension whitelist; otherwise fall back."""
    allowed = ALLOWED_MIME_TYPES[extension]
    if mime_type:
        normalized = mime_type.split(";")[0].strip().lower()
        if normalized in allowed:
            return normalized
        # Browsers sometimes send empty or wrong types — warn but don't hard-fail
        # when the extension is already validated.
        logger.info(
            "MIME type %s does not match extension %s; accepting by extension",
            normalized,
            extension,
        )
    # Canonical MIME for storage when client didn't send a useful one.
    defaults = {
        ".txt": "text/plain",
        ".md": "text/markdown",
        ".markdown": "text/markdown",
        ".csv": "text/csv",
        ".html": "text/html",
        ".htm": "text/html",
        ".rtf": "application/rtf",
        ".pdf": "application/pdf",
        ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    }
    return defaults[extension]


def _decode_text(data: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise FileReaderError("Could not decode file as text")


def _truncate(text: str) -> tuple[str, bool]:
    cleaned = text.replace("\x00", "").strip()
    if len(cleaned) <= MAX_EXTRACTED_CHARS:
        return cleaned, False
    return cleaned[:MAX_EXTRACTED_CHARS].rstrip() + "\n\n[…truncated]", True


class _HTMLTextExtractor(HTMLParser):
    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self._chunks: list[str] = []
        self._skip = False

    def handle_starttag(self, tag: str, attrs) -> None:  # noqa: ANN001
        if tag in {"script", "style", "noscript"}:
            self._skip = True
        elif tag in {"p", "div", "br", "li", "h1", "h2", "h3", "h4", "tr"}:
            self._chunks.append("\n")

    def handle_endtag(self, tag: str) -> None:
        if tag in {"script", "style", "noscript"}:
            self._skip = False
        elif tag in {"p", "div", "li", "h1", "h2", "h3", "h4"}:
            self._chunks.append("\n")

    def handle_data(self, data: str) -> None:
        if not self._skip and data:
            self._chunks.append(data)

    def get_text(self) -> str:
        return re.sub(r"\n{3,}", "\n\n", "".join(self._chunks)).strip()


def _extract_html(data: bytes) -> str:
    parser = _HTMLTextExtractor()
    parser.feed(_decode_text(data))
    parser.close()
    return parser.get_text()


def _extract_rtf(data: bytes) -> str:
    """Minimal RTF → text: strip control words, keep printable runs."""
    raw = _decode_text(data)
    # Remove groups we don't care about (fonts, colors, etc. still leave text).
    text = re.sub(r"\\'[0-9a-fA-F]{2}", " ", raw)
    text = re.sub(r"\\[a-zA-Z]+\-?\d* ?", " ", text)
    text = text.replace("{", " ").replace("}", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _extract_csv(data: bytes) -> str:
    text = _decode_text(data)
    reader = csv.reader(io.StringIO(text))
    rows = []
    for row in reader:
        rows.append(" | ".join(cell.strip() for cell in row))
    return "\n".join(rows)


def _extract_pdf(data: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise FileReaderError(
            "PDF support is not installed on the server", status_code=500
        ) from exc

    try:
        reader = PdfReader(io.BytesIO(data))
    except Exception as exc:
        raise FileReaderError(f"Could not read PDF: {exc}") from exc

    if getattr(reader, "is_encrypted", False):
        try:
            # Empty password unlocks some "owner-restricted" PDFs.
            if reader.decrypt("") == 0:
                raise FileReaderError("PDF is password-protected and cannot be imported")
        except FileReaderError:
            raise
        except Exception as exc:
            raise FileReaderError("PDF is password-protected and cannot be imported") from exc

    parts: list[str] = []
    for page in reader.pages:
        try:
            page_text = page.extract_text() or ""
        except Exception:
            page_text = ""
        if page_text.strip():
            parts.append(page_text.strip())
    return "\n\n".join(parts)


def _extract_docx(data: bytes) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise FileReaderError(
            "DOCX support is not installed on the server", status_code=500
        ) from exc

    try:
        document = Document(io.BytesIO(data))
    except Exception as exc:
        raise FileReaderError(f"Could not read DOCX: {exc}") from exc

    paragraphs = [p.text.strip() for p in document.paragraphs if p.text and p.text.strip()]
    # Also pull simple table cell text so spreadsheets-in-docx aren't lost.
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))
    return "\n\n".join(paragraphs)


def extract_text(
    data: bytes,
    *,
    filename: str,
    mime_type: str | None = None,
) -> ExtractedDocument:
    """Validate and extract text from an uploaded file's bytes."""
    if not data:
        raise FileReaderError("Uploaded file is empty")
    if len(data) > MAX_UPLOAD_BYTES:
        raise FileReaderError(
            f"File exceeds the {MAX_UPLOAD_BYTES // (1024 * 1024)} MB size limit"
        )

    safe_name = sanitize_filename(filename)
    extension = resolve_extension(safe_name)
    resolved_mime = validate_mime(extension, mime_type)

    # Light magic-byte checks for binary formats (extension alone is not enough).
    if extension == ".pdf" and not data.startswith(b"%PDF"):
        raise FileReaderError("File does not look like a valid PDF")
    if extension == ".docx" and not data.startswith(b"PK"):
        raise FileReaderError("File does not look like a valid DOCX (ZIP) archive")

    if extension in {".txt", ".md", ".markdown"}:
        raw_text = _decode_text(data)
    elif extension == ".csv":
        raw_text = _extract_csv(data)
    elif extension in {".html", ".htm"}:
        raw_text = _extract_html(data)
    elif extension == ".rtf":
        raw_text = _extract_rtf(data)
    elif extension == ".pdf":
        raw_text = _extract_pdf(data)
    elif extension == ".docx":
        raw_text = _extract_docx(data)
    else:
        raise FileReaderError(f"Unsupported file type '{extension}'")

    text, truncated = _truncate(raw_text)
    if not text:
        raise FileReaderError(
            "No readable text could be extracted from this file. "
            "Scanned/image-only PDFs are not supported."
        )

    return ExtractedDocument(
        text=text,
        filename=safe_name,
        extension=extension,
        mime_type=resolved_mime,
        truncated=truncated,
    )
